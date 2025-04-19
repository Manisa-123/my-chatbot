import gradio as gr
import requests

from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from langchain.schema import Document
from docx import Document as DocxDocument
import logging

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

gemini_api_key = "AIzaSyCQvlzbSF-26w0JhmNiFblfANKmLuBfry8"
gemini_endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gemini_api_key}"

# Initialize embedding model and ChromaDB client
embedding_model = SentenceTransformer('all-MiniLM-L12-v1')
from chromadb import PersistentClient

text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50, separators=["\n\n", "\n", ".", "!", "?"])
docs_dir = "./documents"
chroma_client = PersistentClient(
    path="./chroma_db/"          # files will live under ./chroma_db/
)


# Load and parse docx files
def load_docx(file_path):
    """Load text from a .docx file."""
    doc = DocxDocument(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return Document(page_content=text)

# Generate embeddings using Sentence-Transformers
def generate_embeddings(text):
    return embedding_model.encode(text)


from chromadb.api.types import EmbeddingFunction


class SentenceTransformerEmbeddingFunction(EmbeddingFunction):
    def __init__(self, model):
        self.model = model

    def __call__(self, input: list) -> list:
        embeddings = self.model.encode(input, show_progress_bar=False)
        return embeddings.tolist()


embedding_function = SentenceTransformerEmbeddingFunction(embedding_model)



def call_gemini_api(query, api_key):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": query}]}]}

    response = requests.post(url, headers=headers, json=payload)
    if response.status_code == 200:
        try:
            data = response.json()
            logging.info(f"Gemini API Response: {data}")
            # if data["candidates"][0]["avgLogprobs"] < -0.1:
            #     return "I don't know"

            generated_text = data.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "Sorry, something went wrong.")
            return generated_text
        except Exception as e:
            logging.info(f"Error parsing Gemini API response: {e}")
            return "Sorry, something went wrong. Please try again."
    else:
        return f"Error from Gemini API: {response.text}"



def get_relevant_docs(query):
    collection = chroma_client.get_or_create_collection(
        name="user_data_1",
        embedding_function=embedding_function
    )
    results = collection.query(
        query_texts=[query],
        n_results=5
    )

    docs_nested = results.get("documents", [])
    # Unwrap the first list (one per query) if needed
    docs_list = docs_nested[0] if docs_nested and isinstance(docs_nested[0], list) else docs_nested

    # Filter out any non‑string or empty values
    texts = [t for t in docs_list if isinstance(t, str) and t.strip()]
    return [Document(page_content=t) for t in texts]




def chatbot(query):
    relevant_docs = get_relevant_docs(query)
    if not relevant_docs:
        return "I DONT know"

    combined_info = "\n".join([doc.page_content for doc in relevant_docs])
    generated_text = call_gemini_api(f"{combined_info}\n\nQuery: {query}", gemini_api_key)


    # normalize to lowercase for matching

    # if ("doesn't contain any information" in lower
    #         or "does not contain any information" in lower
    #         or "doesn't contain information" in lower
    #         or "I do not know" in lower
    #         or "i don't know" in lower
    # or " I cannot assist " in lower
    # or "doesn't contain information" in lower):
    #     return "I DONT know"

    return generated_text


iface = gr.Interface(
    fn=chatbot,
    inputs=gr.Textbox(lines=2, placeholder="Enter your query here..."),
    outputs="text",
    title="Insurance AI Chatbot"
)

if __name__ == "__main__":
    # process_documents()
    iface.launch(share=True)

