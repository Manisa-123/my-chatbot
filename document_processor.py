from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from langchain.schema import Document
from docx import Document as DocxDocument
import os
import hashlib
from scrap import load_webpage_as_document, fetch_support_pages
from chatbot_ui import logger

embedding_model = SentenceTransformer('all-MiniLM-L12-v1')
from chromadb import PersistentClient

text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50, separators=["\n\n", "\n", ".", "!", "?"])
docs_dir = "./documents"
chroma_client = PersistentClient(
    path="./chroma_db/"
)


def load_docx(file_path):
    """Load text from a .docx file."""
    doc = DocxDocument(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return Document(page_content=text)  # Wrap text in LangChain's Document class


def generate_embeddings(text):
    return embedding_model.encode(text)


from chromadb.api.types import EmbeddingFunction


class SentenceTransformerEmbeddingFunction(EmbeddingFunction):
    def __init__(self, model):
        self.model = model

    def __call__(self, input: list) -> list:

        embeddings = self.model.encode(input, show_progress_bar=False)
        return embeddings.tolist()


embedding_function = SentenceTransformerEmbeddingFunction(embedding_model)


def store_embedding(text, embedding, user_id=1):
    # Create or get the collection, passing the embedding function
    collection = chroma_client.get_or_create_collection(
        name=f"user_data_{user_id}",
        embedding_function=embedding_function  # Pass the wrapped embedding function
    )

    doc_id = hashlib.sha256(text.encode()).hexdigest()
    existing_data = collection.get(ids=[doc_id])

    if existing_data and existing_data["ids"]:
        logger.info("Text already exists in ChromaDB, skipping embedding.")
        return doc_id

    collection.add(documents=[text], embeddings=[embedding.tolist()], ids=[doc_id])
    logger.info("New text stored in ChromaDB.")
    return doc_id


def store_informations(text):
    embedding = generate_embeddings(text)
    doc_id = store_embedding(text, embedding)
    return doc_id


def process_documents():
    logger.info("Starting support‑page crawl…")
    documents = []

    # 1) Add all support webpages
    for url in fetch_support_pages():
        logger.info(f"→ Crawling: {url}")
        doc = load_webpage_as_document(url)
        # Only add non‑empty pages
        if doc.page_content.strip():
            documents.extend(text_splitter.split_documents([doc]))


    for file_name in os.listdir(docs_dir):
        file_path = os.path.join(docs_dir, file_name)
        logger.info(f"Processing file: {file_name}")
        if file_name.endswith(".pdf"):
            loader = PyPDFLoader(file_path)
            docs = loader.load()
        elif file_name.endswith(".docx"):
            docs = [load_docx(file_path)]
        else:
            logger.info(f"Skipping unsupported file type: {file_name}")
            continue

        documents.extend(text_splitter.split_documents(docs))

    for doc in documents:
        store_informations(doc.page_content)
    logger.info("All document chunks stored successfully.")


if __name__ == "__main__":
    process_documents()
